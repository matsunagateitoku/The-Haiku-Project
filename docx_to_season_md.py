#!/usr/bin/env python3
"""
Convert a season saijiki .docx (like "Saijiki summer .docx") into a single
season-level Markdown file under saijiki_seasons/.

The source docx contains, for many kigo, TWO drafts: an early terse pass
(Heading 1 / Heading 2 styles, bare poems, no translation split) and a later
fuller pass (Title style, prose + variations + translated poems). Where both
exist, this script keeps the LATER (fuller) occurrence, since spot checks
confirm entries already extracted into essays/*.txt match the fuller version.
Entries that only ever received the terse pass are kept in terse form.

This script deliberately does not invent structure the source doesn't have:
category dividers (The Heavens / The Earth / Daily Life / Observances) are
preserved only where literally present; no headings are synthesized.
"""
import re
import sys
import docx

CJK_RE = re.compile(r'[぀-ヿ㐀-鿿ｦ-ﾟ]')
KNOWN_CATEGORY_TITLES = {"The Season", "The Heavens", "The Earth", "Daily Life", "Observances"}
SKIP_TITLES = {"Table of Contents"}
TAB_JUNK_RE = re.compile(r'^Tab \d+$')

def leading_cjk_run(text):
    m = re.match(r'\s*([぀-ヿ㐀-鿿ｦ-ﾟ]+)', text)
    return m.group(1) if m else None

def normalize_ws(s):
    # collapse any whitespace run containing a tab or ideographic space
    # (the source's field separator between poem/romaji text and poet name)
    # into " -- "; leave ordinary single spaces alone.
    return re.sub(r'[ \t　]*[\t　][ \t　]*', ' -- ', s.strip())

def is_blank(p):
    return not p.text.strip()

def related_titles(a, b):
    ka, kb = leading_cjk_run(a), leading_cjk_run(b)
    if ka and kb and ka == kb:
        return True
    al, bl = a.lower(), b.lower()
    return al in bl or bl in al

def entry_key(title_text, body_paragraphs):
    key = leading_cjk_run(title_text)
    if key:
        return key
    for p in body_paragraphs[:6]:
        t = p.text.strip()
        if not t or t == title_text.strip():
            continue
        k = leading_cjk_run(t)
        if k:
            return k
    return title_text.strip()

NUM_PREFIX_RE = re.compile(r'^\d+\.\s*')

def starts_with_cjk(text):
    return bool(CJK_RE.match(NUM_PREFIX_RE.sub('', text.strip())))

def parse_docx(path):
    d = docx.Document(path)
    paras = d.paragraphs
    heading_idxs = [i for i, p in enumerate(paras)
                    if p.style.name in ("Title", "Heading 1")]
    heading_idxs.append(len(paras))  # sentinel

    # Merge runs of consecutive headings that restate the same entry with no
    # content of their own (e.g. a Title immediately followed by a Heading 1
    # repeating the same kigo, with the real body only starting after the
    # last one). Only merge when the titles are actually related, so a truly
    # empty stub heading (e.g. an unwritten entry) followed by an unrelated
    # entry is left alone as its own empty stub.
    merged_bounds = []  # list of (title_index_to_use, body_start, body_end)
    n = 0
    while n < len(heading_idxs) - 1:
        run_titles = [paras[heading_idxs[n]].text.strip()]
        while True:
            i, j = heading_idxs[n], heading_idxs[n + 1]
            body_has_content = any(paras[k].text.strip() for k in range(i + 1, j))
            if body_has_content or n + 1 >= len(heading_idxs) - 1:
                break
            next_title = paras[heading_idxs[n + 1]].text.strip()
            if not any(related_titles(next_title, rt) for rt in run_titles):
                break
            run_titles.append(next_title)
            n += 1
        title_idx = heading_idxs[n]
        body_start = heading_idxs[n] + 1
        body_end = heading_idxs[n + 1]
        merged_bounds.append((title_idx, body_start, body_end))
        n += 1

    raw_entries = []  # list of dicts: index, title, style, body(list of para objects)
    dropped_junk = []
    for title_idx, body_start, body_end in merged_bounds:
        title_para = paras[title_idx]
        title_text = title_para.text.strip()
        body = [paras[k] for k in range(body_start, body_end)]

        if title_text in SKIP_TITLES:
            continue
        if TAB_JUNK_RE.match(title_text):
            dropped_junk.append((title_idx, title_text))
            continue

        raw_entries.append({
            "index": title_idx,
            "title": title_text,
            "style": title_para.style.name,
            "body": body,
        })

    return raw_entries, dropped_junk

def strip_dup_title_lines(title_text, body):
    out = []
    for p in body:
        t = p.text.strip()
        if t == title_text:
            continue
        out.append(p)
    return out

def find_poem_start(body):
    """Return index into body (list of paragraphs) where the poems section begins."""
    for i, p in enumerate(body):
        t = p.text.strip()
        if not t:
            continue
        if p.style.name == "Heading 2":
            return i
        if starts_with_cjk(t):
            # look ahead to next non-blank paragraph
            for k in range(i + 1, len(body)):
                nt = body[k].text.strip()
                if not nt:
                    continue
                if '/' in nt and not starts_with_cjk(nt):
                    return i
                break
    return len(body)

def group_terse_poems(poem_body):
    groups = []
    i = 0
    while i < len(poem_body):
        p = poem_body[i]
        if p.style.name == "Heading 2" and p.text.strip():
            jp = re.sub(r'^\s*\d+\.\s*', '', p.text.strip())
            rest = []
            k = i + 1
            while k < len(poem_body) and poem_body[k].style.name != "Heading 2":
                if poem_body[k].text.strip():
                    rest.append(poem_body[k].text.strip())
                k += 1
            groups.append({"jp": jp, "rest": rest})
            i = k
        else:
            i += 1
    return groups

def group_rich_poems(poem_body):
    texts = [p.text.strip() for p in poem_body]
    groups = []
    i = 0
    n = len(texts)
    # find all indices that are JP-poem-line starts
    starts = []
    for i in range(n):
        t = texts[i]
        if not t or not starts_with_cjk(t):
            continue
        for k in range(i + 1, n):
            if not texts[k]:
                continue
            if '/' in texts[k] and not starts_with_cjk(texts[k]):
                starts.append(i)
            break
    starts.append(n)
    for s in range(len(starts) - 1):
        a, b = starts[s], starts[s + 1]
        block = [t for t in texts[a:b] if t]
        if not block:
            continue
        jp = NUM_PREFIX_RE.sub('', block[0])
        romaji = block[1] if len(block) > 1 else ""
        translation_lines = block[2:]
        groups.append({"jp": jp, "romaji": romaji, "translation": translation_lines})
    return groups

def render_entry_md(entry, warnings):
    title = entry["title"]
    body = strip_dup_title_lines(title, entry["body"])
    # drop fully blank entry (stub)
    non_blank = [p for p in body if p.text.strip()]
    if not non_blank:
        return f"### {title}\n\n_(stub in source -- no content drafted yet)_\n"

    p_start = find_poem_start(body)
    front = [p.text.strip() for p in body[:p_start] if p.text.strip()]
    poem_body = body[p_start:]

    lines = [f"### {title}", ""]
    for para_text in front:
        # preserve internal newlines (python-docx line breaks) as separate lines
        for sub in para_text.split("\n"):
            sub = sub.strip()
            if sub:
                lines.append(normalize_ws(sub) if leading_cjk_run(sub) else sub)
        lines.append("")

    if poem_body and poem_body[0].style.name == "Heading 2":
        groups = group_terse_poems(poem_body)
        fmt = "terse"
    else:
        groups = group_rich_poems(poem_body)
        fmt = "rich"

    if not groups:
        warnings.append(f"[{title}] no poems detected")
    else:
        lines.append("**Poems:**")
        lines.append("")
        for idx, g in enumerate(groups, 1):
            lines.append(f"{idx}. {normalize_ws(g['jp'])}")
            if fmt == "rich":
                if g["romaji"]:
                    lines.append(f"   *{normalize_ws(g['romaji'])}*")
                for tl in g["translation"]:
                    lines.append(f"   > {tl}")
            else:
                for tl in g["rest"]:
                    lines.append(f"   {tl}")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"

def convert(path, season_label):
    raw_entries, dropped_junk = parse_docx(path)

    # dedup: keep LAST occurrence per key among non-category, non-stub entries
    keyed = {}
    order = []
    category_positions = []  # (index, title) for category dividers, in doc order
    stubs = []

    for e in raw_entries:
        title = e["title"]
        body = strip_dup_title_lines(title, e["body"])
        non_blank = [p for p in body if p.text.strip()]

        if title in KNOWN_CATEGORY_TITLES:
            category_positions.append((e["index"], title))
            continue

        if title == season_label or (e is raw_entries[0] and not non_blank):
            pass

        if not non_blank:
            stubs.append((e["index"], title))
            keyed_entry_key = title  # unique, won't collide
            e["_key"] = keyed_entry_key
        else:
            e["_key"] = entry_key(title, body)

        if e["_key"] in keyed:
            order.remove(keyed[e["_key"]]["index"])
        keyed[e["_key"]] = e
        order.append(e["index"])

    order.sort()
    idx_to_entry = {e["index"]: e for e in keyed.values()}
    winning_entries = [idx_to_entry[i] for i in order]

    # Drop redundant empty stubs whose title merely restates a real
    # (non-stub) winning entry elsewhere in the document.
    stub_indices = {i for i, _ in stubs}
    real_titles = [e["title"] for e in winning_entries if e["index"] not in stub_indices]
    winning_entries = [
        e for e in winning_entries
        if e["index"] not in stub_indices
        or not any(related_titles(e["title"], rt) for rt in real_titles)
    ]

    # merge categories + entries into one document-order stream
    stream = [("cat", i, t) for i, t in category_positions] + \
             [("entry", e["index"], e) for e in winning_entries]
    stream.sort(key=lambda x: x[1])

    warnings = []
    out = []
    for kind, i, payload in stream:
        if kind == "cat":
            out.append(f"## {payload}\n")
        else:
            out.append(render_entry_md(payload, warnings))

    return "\n".join(out), warnings, dropped_junk, stubs, len(winning_entries)

if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "resources/Saijiki summer .docx"
    out_path = sys.argv[2] if len(sys.argv) > 2 else "saijiki_seasons/summer.md"

    d = docx.Document(src)
    # season title + intro is the first Title paragraph with real content
    season_title_idx = next(i for i, p in enumerate(d.paragraphs)
                             if p.style.name == "Title" and p.text.strip() not in SKIP_TITLES | KNOWN_CATEGORY_TITLES)
    season_label = d.paragraphs[season_title_idx].text.strip()

    body_md, warnings, dropped_junk, stubs, n_entries = convert(src, season_label)

    header = f"# {season_label}\n\n"
    full_md = header + body_md

    import os
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(full_md)

    print(f"Wrote {out_path}: {n_entries} entries")
    print(f"Dropped junk titles ({len(dropped_junk)}): {[t for _, t in dropped_junk]}")
    print(f"Stub (empty) entries ({len(stubs)}): {[t for _, t in stubs]}")
    if warnings:
        print(f"Warnings ({len(warnings)}):")
        for w in warnings:
            print(" -", w)
